import os
import streamlit as st
import speech_recognition as sr
from docx import Document
from pydub import AudioSegment
import tempfile
import time
import io

# Uygulama başlığı ve açıklaması
st.set_page_config(page_title="Ses Transkript Uygulaması", layout="wide")
st.title("Ses Transkript Uygulaması")
st.markdown("Ses dosyanızı yükleyin, transkript oluşturun ve Word belgesi olarak indirin.")

def convert_audio(audio_bytes, file_extension):
    """Farklı ses formatlarını speech_recognition ile işlenebilecek formata dönüştürür."""
    try:
        # Geçici dosya oluştur (orijinal uzantıyla)
        temp_input_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}')
        temp_input_path = temp_input_file.name
        temp_input_file.close()
        
        # Geçici WAV dosyası oluştur
        temp_output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        temp_output_path = temp_output_file.name
        temp_output_file.close()
        
        # Byte verisini dosyaya yaz
        with open(temp_input_path, 'wb') as f:
            f.write(audio_bytes)
        
        # Ses formatını WAV'a dönüştür
        audio = AudioSegment.from_file(temp_input_path)
        audio.export(temp_output_path, format="wav")
        
        return temp_output_path, audio, temp_input_path
    except Exception as e:
        st.error(f"Dönüştürme hatası: {e}")
        return None, None, None

def split_audio(audio, chunk_length_ms=50000):
    """Ses dosyasını belirtilen uzunlukta parçalara böler."""
    chunks = []
    
    # Ses dosyasını parçalara ayır
    for i in range(0, len(audio), chunk_length_ms):
        chunk = audio[i:i+chunk_length_ms]
        chunks.append(chunk)
    
    return chunks

def transcribe_audio(audio_path):
    """Ses dosyasını metne çevirir."""
    r = sr.Recognizer()
    
    try:
        with sr.AudioFile(audio_path) as source:
            audio_data = r.record(source)
            text = r.recognize_google(audio_data, language="tr-TR")  # Türkçe tanıma
            return text
    except sr.UnknownValueError:
        return "[Anlaşılamayan Bölüm]"
    except sr.RequestError as e:
        return f"[Servis Hatası: {e}]"
    except Exception as e:
        return f"[Hata: {e}]"

def create_docx(transcript_text):
    """Transkripti Word belgesine dönüştürür."""
    document = Document()
    document.add_heading("Ses Transkripsiyonu", 0)
    document.add_paragraph(transcript_text)
    
    # Dosyayı belleğe kaydet
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def process_audio(audio_file):
    """Ses dosyasını işleyip transkript eder"""
    if audio_file is None:
        return None, "Lütfen bir ses dosyası yükleyin."
    
    # Progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Dosya uzantısını al
    file_extension = audio_file.name.split('.')[-1].lower()
    
    # Ses dosyasını oku ve dönüştür
    status_text.text("Ses dosyası dönüştürülüyor...")
    progress_bar.progress(10)
    
    wav_path, audio, temp_input_path = convert_audio(audio_file.read(), file_extension)
    
    if wav_path is None:
        return None, "Ses dosyası dönüştürülemedi."
    
    try:
        # Ses dosyasını parçalara böl
        status_text.text("Ses parçalara ayrılıyor...")
        progress_bar.progress(20)
        
        chunks = split_audio(audio)
        full_transcript = ""
        
        # Her parçayı transkript et
        for i, chunk in enumerate(chunks):
            progress_value = 20 + int(60 * (i + 1) / len(chunks))
            status_text.text(f"Parça {i+1}/{len(chunks)} transkript ediliyor...")
            progress_bar.progress(progress_value)
            
            # Geçici chunk dosyasını oluştur
            chunk_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
            chunk_path = chunk_file.name
            chunk_file.close()
            
            chunk.export(chunk_path, format="wav")
            
            # Parçayı transkript et
            chunk_text = transcribe_audio(chunk_path)
            full_transcript += chunk_text + " "
            
            # Geçici chunk dosyasını temizle
            os.remove(chunk_path)
        
        # Word belgesi oluştur
        status_text.text("Word belgesi oluşturuluyor...")
        progress_bar.progress(90)
        
        docx_bytes = create_docx(full_transcript.strip())
        
        # Temizlik
        status_text.text("İşlem tamamlandı!")
        progress_bar.progress(100)
        
        return docx_bytes, full_transcript.strip()
    except Exception as e:
        return None, f"Hata oluştu: {e}"
    finally:
        # Geçici dosyaları temizle
        for path in [wav_path, temp_input_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except:
                    pass

# Ses dosyası yükleme alanı
uploaded_file = st.file_uploader("Ses dosyası yükleyin", type=["wav", "mp3", "ogg", "m4a", "mp4"])

# İşlem butonları
col1, col2 = st.columns(2)
with col1:
    process_button = st.button("Transkript Et", type="primary")

# Eğer dosya yüklendi ve buton tıklandıysa
if uploaded_file is not None and process_button:
    docx_bytes, transcript = process_audio(uploaded_file)
    
    if docx_bytes:
        st.markdown("### Transkript Metin")
        st.write(transcript)
        
        st.markdown("### Word Belgesi")
        st.download_button(
            label="Word Belgesini İndir",
            data=docx_bytes,
            file_name="transkript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error(transcript)  # Hata mesajını göster

# Kullanım talimatları
st.markdown("---")
st.markdown("### Nasıl Kullanılır")
st.markdown("""
1. Bir ses dosyası yükleyin (WAV, MP3, OGG, M4A, MP4 formatları desteklenir)
2. "Transkript Et" düğmesine tıklayın
3. Transkript çıkarıldıktan sonra, Word belgesini indirebilirsiniz

**Not:** Uzun kayıtlar için işlem biraz zaman alabilir. Lütfen bekleyin.
""") 